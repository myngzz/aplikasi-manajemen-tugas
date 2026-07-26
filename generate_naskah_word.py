from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

output_path = 'Naskah_Presentasi_Aplikasi_Manajemen_Tugas_Sederhana.docx'

doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal'].font.size = Pt(11)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Naskah Presentasi Video\nAplikasi Manajemen Tugas Sederhana')
run.bold = True
run.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Durasi target: maksimal 10 menit')
run.italic = True
run.font.size = Pt(11)

doc.add_paragraph('')
intro = doc.add_paragraph()
intro.add_run('Catatan: ').bold = True
intro.add_run('Naskah ini bisa dibacakan langsung saat merekam video presentasi. Bagian dalam tanda kurung boleh dibaca seperlunya atau dijadikan arahan saat demo layar.')

sections = [
    ('0:00 - 0:40 | Pembukaan', [
        'Assalamualaikum warahmatullahi wabarakatuh.',
        'Perkenalkan, saya akan mempresentasikan aplikasi web sederhana yang menerapkan sistem terdistribusi client-server dengan tiga peran, yaitu admin, user, dan server sebagai perantara.',
        'Aplikasi ini dibuat untuk menunjukkan bagaimana tugas dikirim dari admin, diterima oleh user, lalu statusnya diperbarui melalui server yang sama.'
    ]),
    ('0:40 - 1:20 | Judul dan tujuan', [
        'Nama aplikasi ini adalah Aplikasi Manajemen Tugas Sederhana.',
        'Tujuan aplikasi ini adalah membantu proses pembagian tugas dari admin ke user secara terstruktur, lalu memudahkan user untuk melihat, menerima, dan menyelesaikan tugas tersebut.',
        'Sistem ini juga menunjukkan konsep sistem terdistribusi sederhana karena ada pemisahan peran antara client dan server, serta komunikasi data melalui API.'
    ]),
    ('1:20 - 2:00 | Penjelasan konsep', [
        'Di aplikasi ini ada tiga bagian utama.',
        'Pertama, halaman admin untuk membuat dan mengirim tugas.',
        'Kedua, halaman user untuk menerima dan menyelesaikan tugas.',
        'Ketiga, dashboard server untuk memantau data dan status API.',
        'Server di sini tidak hanya menyimpan data, tetapi juga menjadi perantara antara admin dan user.'
    ]),
    ('2:00 - 3:15 | Demo halaman admin', [
        'Sekarang saya masuk ke halaman admin.',
        'Di halaman ini, admin bisa membuat tugas baru dengan mengisi judul tugas, nama user penerima, deskripsi, dan prioritas.',
        'Sebagai contoh, saya akan memberikan tugas kepada user bernama Budi.',
        'Setelah data diisi, admin menekan tombol simpan tugas.',
        'Setelah tugas dikirim, data langsung masuk ke server dan disimpan di storage JSON.',
        'Di bagian bawah juga terlihat daftar tugas yang sudah tersimpan beserta statusnya.'
    ]),
    ('3:15 - 4:30 | Demo halaman user', [
        'Selanjutnya saya buka halaman user.',
        'Di sini user bisa memilih namanya terlebih dahulu, misalnya Budi.',
        'Setelah nama dipilih, sistem akan menampilkan tugas yang memang dikirim untuk user tersebut.',
        'User kemudian bisa menekan tombol Terima Tugas untuk mengubah status dari ditugaskan menjadi diterima.',
        'Setelah itu, user juga bisa menekan tombol Selesai untuk menandai bahwa tugas telah selesai dikerjakan.',
        'Jadi, halaman user berfungsi untuk menerima dan menyelesaikan tugas yang berasal dari admin.'
    ]),
    ('4:30 - 5:45 | Demo dashboard server', [
        'Sekarang saya buka dashboard server.',
        'Halaman ini berfungsi sebagai pusat pemantauan.',
        'Di sini terlihat status API, jumlah tugas, jumlah tugas yang sedang diproses, dan jumlah tugas yang sudah selesai.',
        'Dashboard server juga menampilkan ringkasan request seperti GET, POST, PATCH, dan DELETE.',
        'Ini menunjukkan bahwa server benar-benar bekerja sebagai penghubung antara admin dan user.'
    ]),
    ('5:45 - 6:45 | Penjelasan alur data', [
        'Alur datanya seperti ini.',
        'Admin mengirim tugas ke server melalui halaman admin.',
        'Server menyimpan data ke file JSON.',
        'User mengambil tugas dari server melalui halaman user.',
        'Saat user menerima atau menyelesaikan tugas, statusnya dikirim kembali ke server dan disimpan ulang.',
        'Dengan begitu, data tetap sinkron antara halaman admin, halaman user, dan dashboard server.'
    ]),
    ('6:45 - 7:30 | Kelebihan aplikasi', [
        'Kelebihan dari aplikasi ini adalah alurnya sederhana, mudah dipahami, dan cocok untuk mempresentasikan konsep client-server.',
        'Selain itu, data tersimpan secara persisten, jadi walaupun halaman direfresh, tugas tetap ada.',
        'Aplikasi ini juga ringan, mudah dijalankan, dan bisa menjadi contoh dasar untuk sistem terdistribusi yang lebih kompleks.'
    ]),
    ('7:30 - 8:10 | Penutup', [
        'Kesimpulannya, aplikasi ini berhasil menunjukkan konsep sistem terdistribusi sederhana dengan pembagian peran admin, user, dan server.',
        'Admin bertugas memberi tugas, user menerima dan menyelesaikan tugas, sedangkan server menjadi perantara penyimpanan dan pertukaran data.',
        'Terima kasih atas perhatiannya.',
        'Wassalamualaikum warahmatullahi wabarakatuh.'
    ])
]

for heading, lines in sections:
    doc.add_heading(heading, level=1)
    for line in lines:
        doc.add_paragraph(line)

# Closing note

doc.add_page_break()
doc.add_heading('Catatan Rekaman', level=1)
doc.add_paragraph('• Jika perlu, bacakan dengan tempo normal dan beri jeda saat pindah dari admin ke user dan server.')
doc.add_paragraph('• Saat demo, tampilkan alur: admin mengirim tugas, user menerima, lalu server memperbarui status.')
doc.add_paragraph('• Total durasi aman jika dibacakan sekitar 7 sampai 9 menit.')

doc.save(output_path)
print(f'Saved to {output_path}')
